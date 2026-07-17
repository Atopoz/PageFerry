"""验证 DOCX 的 hyperlink、SDT、字段与 note XML 结构边界。"""

import os
import zipfile
from collections.abc import Callable, Sequence
from pathlib import Path

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree

from modules.docx import DocxPipeline
from modules.docx.extractor import DocxExtractor
from modules.translation.contracts import TranslationBatchItem, TranslationBatchResult


class StructureTranslator:
    """替换指定文本并完整保留结构 marker。"""

    def __init__(self, transform: Callable[[str], str]) -> None:
        """保存结构测试使用的确定性转换函数。"""

        self._transform = transform

    def translate_batch(
        self,
        *,
        texts: Sequence[str],
        source_language: str | None,
        target_language: str,
        format_hint: str,
        read_only_context: Sequence[str] = (),
    ) -> TranslationBatchResult:
        """返回与输入 index 一一对应的结构安全结果。"""

        del source_language, target_language, format_hint, read_only_context
        return TranslationBatchResult(
            items=tuple(
                TranslationBatchItem(index=index, text=self._transform(text))
                for index, text in enumerate(texts)
            )
        )


def test_hyperlink_sdt_and_field_boundaries_survive_translation(tmp_path: Path) -> None:
    """不同 XML 容器和字段两侧 run 不合并, 且回填后节点仍在原位。"""

    source = tmp_path / "structures.docx"
    output = tmp_path / "translated.docx"
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("A")
    _append_hyperlink(paragraph, "B", "https://example.com")
    _append_inline_sdt(paragraph, "C")
    _append_field(paragraph, " PAGE ")
    paragraph.add_run("D")
    document.save(source)

    extracted = DocxExtractor().extract(Document(source), source)
    assert len(extracted) == 1
    assert [run.text for run in extracted[0].original_runs] == ["A", "B", "C", "D"]

    translator = StructureTranslator(
        lambda text: (
            text.replace(">A<", ">甲<")
            .replace(">B<", ">乙<")
            .replace(">C<", ">丙<")
            .replace(">D<", ">丁<")
        )
    )
    DocxPipeline(translator, translate_tables=False).translate_to(
        source_path=source,
        output_path=output,
        source_language="en",
        target_language="zh-CN",
    )

    with zipfile.ZipFile(output) as package:
        root = etree.fromstring(package.read("word/document.xml"))
    assert "".join(root.itertext()).replace(" PAGE ", "") == "甲乙丙丁"
    assert len(root.findall(".//" + qn("w:hyperlink"))) == 1
    assert len(root.findall(".//" + qn("w:sdt"))) == 1
    assert len(root.findall(".//" + qn("w:fldChar"))) == 2
    assert [node.text for node in root.findall(".//" + qn("w:instrText"))] == [" PAGE "]


def test_text_after_tab_in_same_run_is_translated_without_moving_tab(tmp_path: Path) -> None:
    """mixed run 中的 w:t 应翻译, 同级 w:tab 仍留在文字之前。"""

    source = tmp_path / "tab-heading.docx"
    output = tmp_path / "translated.docx"
    document = Document()
    paragraph = document.add_paragraph()
    number_run = paragraph.add_run("3.")
    number_run.bold = True
    heading_run = paragraph.add_run()
    heading_run.bold = True
    heading_run._r.append(OxmlElement("w:tab"))
    heading_text = OxmlElement("w:t")
    heading_text.text = "Receipt matters"
    heading_run._r.append(heading_text)
    document.save(source)

    extracted = DocxExtractor().extract(Document(source), source)
    assert [run.text for run in extracted[0].original_runs] == ["3.", "Receipt matters"]
    assert extracted[0].batch_token_text == "[PARA_0]<span>3.\tReceipt matters</span>"

    DocxPipeline(
        StructureTranslator(lambda text: text.replace("Receipt matters", "收据事宜")),
        translate_tables=False,
    ).translate_to(
        source_path=source,
        output_path=output,
        source_language="en",
        target_language="zh-CN",
    )

    translated_paragraph = Document(output).paragraphs[0]
    assert translated_paragraph.text == "3.\t收据事宜"
    translated_mixed_run = translated_paragraph.runs[1]._r
    child_tags = [etree.QName(child).localname for child in translated_mixed_run]
    assert child_tags[-2:] == ["tab", "t"]
    assert len(translated_mixed_run.findall(qn("w:tab"))) == 1


def test_footnote_and_endnote_are_translated_without_losing_references(
    tmp_path: Path,
) -> None:
    """翻译 note part 文本并保留 document relationship 与引用节点。"""

    source = tmp_path / "notes.docx"
    output = tmp_path / "translated.docx"
    document = Document()
    document.add_paragraph("Body text")
    document.save(source)
    _add_note_parts(source)
    original_bytes = source.read_bytes()

    translator = StructureTranslator(
        lambda text: (
            text.replace("Body text", "正文")
            .replace("Footnote text", "脚注")
            .replace("Endnote text", "尾注")
        )
    )
    result = DocxPipeline(translator, translate_tables=False).translate_to(
        source_path=source,
        output_path=output,
        source_language="en",
        target_language="zh-CN",
    )

    assert source.read_bytes() == original_bytes
    assert Document(output).paragraphs[0].text == "正文"
    with zipfile.ZipFile(output) as package:
        document_root = etree.fromstring(package.read("word/document.xml"))
        footnotes = etree.fromstring(package.read("word/footnotes.xml"))
        endnotes = etree.fromstring(package.read("word/endnotes.xml"))
        relationships = etree.fromstring(package.read("word/_rels/document.xml.rels"))
    assert _note_text(footnotes, "footnote", "1") == "脚注"
    assert _note_text(endnotes, "endnote", "1") == "尾注"
    assert len(document_root.findall(".//" + qn("w:footnoteReference"))) == 1
    assert len(document_root.findall(".//" + qn("w:endnoteReference"))) == 1
    relationship_types = {relationship.get("Type") for relationship in relationships}
    assert f"{RT.FOOTNOTES}" in relationship_types
    assert f"{RT.ENDNOTES}" in relationship_types
    assert result.translated_segments == 3
    assert result.fallback_segments == 0


def test_tables_below_depth_limit_are_left_unchanged(tmp_path: Path) -> None:
    """翻译 depth 0..3 的表格, 并让第四层及更深子表原样保留。"""

    source = tmp_path / "nested-tables.docx"
    output = tmp_path / "translated.docx"
    document = Document()
    current_table = document.add_table(rows=1, cols=1)
    for level in range(5):
        current_table.cell(0, 0).text = f"Level {level}"
        if level < 4:
            current_table = current_table.cell(0, 0).add_table(rows=1, cols=1)
    document.save(source)
    original_bytes = source.read_bytes()

    translator = StructureTranslator(
        lambda text: _replace_nested_table_levels(text, translated_levels=range(5))
    )
    result = DocxPipeline(translator).translate_to(
        source_path=source,
        output_path=output,
        source_language="en",
        target_language="zh-CN",
    )

    assert source.read_bytes() == original_bytes
    translated = Document(output)
    current_table = translated.tables[0]
    for level in range(5):
        expected = f"层级 {level}" if level <= 3 else f"Level {level}"
        assert current_table.cell(0, 0).paragraphs[0].text == expected
        if level < 4:
            current_table = current_table.cell(0, 0).tables[0]
    assert result.translated_segments == 4
    assert result.fallback_segments == 0


def _append_hyperlink(paragraph, text: str, url: str) -> None:
    """向测试段落追加一个外部 hyperlink run。"""

    relationship_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _replace_nested_table_levels(text: str, *, translated_levels: range) -> str:
    """把指定层级的测试文本替换为可识别译文。"""

    for level in translated_levels:
        text = text.replace(f"Level {level}", f"层级 {level}")
    return text


def _append_inline_sdt(paragraph, text: str) -> None:
    """向测试段落追加一个 inline 内容控件。"""

    control = OxmlElement("w:sdt")
    content = OxmlElement("w:sdtContent")
    run = OxmlElement("w:r")
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    content.append(run)
    control.append(content)
    paragraph._p.append(control)


def _append_field(paragraph, instruction: str) -> None:
    """追加 begin/instruction/end 字段节点以验证结构边界。"""

    begin_run = OxmlElement("w:r")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin_run.append(begin)
    instruction_run = OxmlElement("w:r")
    instruction_element = OxmlElement("w:instrText")
    instruction_element.set(qn("xml:space"), "preserve")
    instruction_element.text = instruction
    instruction_run.append(instruction_element)
    end_run = OxmlElement("w:r")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run.append(end)
    paragraph._p.extend((begin_run, instruction_run, end_run))


def _add_note_parts(path: Path) -> None:
    """给最小 DOCX package 加入可被 python-docx 原样保留的 note parts。"""

    with zipfile.ZipFile(path) as source:
        entries = {info.filename: source.read(info) for info in source.infolist()}
        infos = {info.filename: info for info in source.infolist()}

    entries["[Content_Types].xml"] = _add_note_content_types(entries["[Content_Types].xml"])
    entries["word/_rels/document.xml.rels"] = _add_note_relationships(
        entries["word/_rels/document.xml.rels"]
    )
    entries["word/document.xml"] = _add_note_references(entries["word/document.xml"])
    entries["word/footnotes.xml"] = _note_part_xml("footnote", "Footnote text")
    entries["word/endnotes.xml"] = _note_part_xml("endnote", "Endnote text")

    temporary = path.with_name(f".{path.name}.notes.tmp")
    with zipfile.ZipFile(temporary, "w", compression=zipfile.ZIP_DEFLATED) as target:
        for name, data in entries.items():
            info = infos.get(name)
            target.writestr(info if info is not None else name, data)
    os.replace(temporary, path)


def _add_note_content_types(content: bytes) -> bytes:
    """给 Content_Types 增加 footnote 与 endnote Override。"""

    root = etree.fromstring(content)
    namespace = "http://schemas.openxmlformats.org/package/2006/content-types"
    for note_kind in ("footnotes", "endnotes"):
        override = etree.SubElement(root, f"{{{namespace}}}Override")
        override.set("PartName", f"/word/{note_kind}.xml")
        override.set(
            "ContentType",
            f"application/vnd.openxmlformats-officedocument.wordprocessingml.{note_kind}+xml",
        )
    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)


def _add_note_relationships(content: bytes) -> bytes:
    """给 document relationship 增加两个 note target。"""

    root = etree.fromstring(content)
    namespace = "http://schemas.openxmlformats.org/package/2006/relationships"
    for index, note_kind in enumerate(("footnotes", "endnotes"), start=900):
        relationship = etree.SubElement(root, f"{{{namespace}}}Relationship")
        relationship.set("Id", f"rId{index}")
        relationship.set(
            "Type",
            f"http://schemas.openxmlformats.org/officeDocument/2006/relationships/{note_kind}",
        )
        relationship.set("Target", f"{note_kind}.xml")
    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)


def _add_note_references(content: bytes) -> bytes:
    """在正文末尾添加结构型 footnote/endnote 引用 run。"""

    root = etree.fromstring(content)
    paragraph = root.find(".//" + qn("w:p"))
    assert paragraph is not None
    for reference_name in ("footnoteReference", "endnoteReference"):
        run = etree.SubElement(paragraph, qn("w:r"))
        reference = etree.SubElement(run, qn(f"w:{reference_name}"))
        reference.set(qn("w:id"), "1")
    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)


def _note_part_xml(note_kind: str, text: str) -> bytes:
    """生成包含 separator 与一条正文 note 的最小 XML part。"""

    namespace = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    root = etree.Element(qn(f"w:{note_kind}s"), nsmap={"w": namespace})
    separator = etree.SubElement(root, qn(f"w:{note_kind}"))
    separator.set(qn("w:type"), "separator")
    separator.set(qn("w:id"), "-1")
    separator_paragraph = etree.SubElement(separator, qn("w:p"))
    separator_run = etree.SubElement(separator_paragraph, qn("w:r"))
    etree.SubElement(separator_run, qn("w:separator"))
    note = etree.SubElement(root, qn(f"w:{note_kind}"))
    note.set(qn("w:id"), "1")
    paragraph = etree.SubElement(note, qn("w:p"))
    run = etree.SubElement(paragraph, qn("w:r"))
    text_element = etree.SubElement(run, qn("w:t"))
    text_element.text = text
    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)


def _note_text(root, note_kind: str, note_id: str) -> str:
    """读取指定 note id 的全部 w:t 文本。"""

    for note in root.findall(qn(f"w:{note_kind}")):
        if note.get(qn("w:id")) == note_id:
            return "".join(node.text or "" for node in note.iter(qn("w:t")))
    raise AssertionError(f"missing {note_kind} {note_id}")
