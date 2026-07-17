"""验证 DOCX pipeline 的翻译、修复、回退与原子发布边界。"""

from collections.abc import Callable, Sequence
from pathlib import Path

import pytest
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from modules.docx import DocxPipeline
from modules.docx.pipeline import _package_signature
from modules.translation.contracts import (
    TranslationBatchItem,
    TranslationBatchResult,
    TranslationProgress,
    TranslationRequest,
)


class RecordingTranslator:
    """用确定性函数翻译并记录每次 format hint。"""

    def __init__(self, transform: Callable[[str], str]) -> None:
        """保存用于每个 batch item 的转换函数。"""

        self._transform = transform
        self.calls: list[tuple[str, tuple[str, ...]]] = []

    def translate_batch(
        self,
        *,
        texts: Sequence[str],
        source_language: str | None,
        target_language: str,
        format_hint: str,
        read_only_context: Sequence[str] = (),
        repair_candidates: Sequence[str] = (),
    ) -> TranslationBatchResult:
        """返回 index 与输入严格对齐的转换结果。"""

        del source_language, target_language, read_only_context, repair_candidates
        captured = tuple(texts)
        self.calls.append((format_hint, captured))
        return TranslationBatchResult(
            items=tuple(
                TranslationBatchItem(index=index, text=self._transform(text))
                for index, text in enumerate(captured)
            )
        )


class RepairingTranslator:
    """首次破坏 span, 修复调用时返回合法译文。"""

    def __init__(self) -> None:
        """记录正常调用与唯一一次修复调用。"""

        self.hints: list[str] = []
        self.repair_candidates: list[tuple[str, ...]] = []

    def translate_batch(
        self,
        *,
        texts: Sequence[str],
        source_language: str | None,
        target_language: str,
        format_hint: str,
        read_only_context: Sequence[str] = (),
        repair_candidates: Sequence[str] = (),
    ) -> TranslationBatchResult:
        """根据 format hint 返回损坏或已修复内容。"""

        del source_language, target_language
        self.hints.append(format_hint)
        del read_only_context
        self.repair_candidates.append(tuple(repair_candidates))
        values = (
            tuple(text.replace("Original", "Translated") for text in texts)
            if format_hint.endswith("_repair")
            else tuple("<span>broken" for _text in texts)
        )
        return TranslationBatchResult(
            items=tuple(
                TranslationBatchItem(index=index, text=text) for index, text in enumerate(values)
            )
        )


class FailingTranslator:
    """每次都抛错以稳定触发源文回退。"""

    def __init__(self) -> None:
        """记录 provider 被调用的 hint。"""

        self.hints: list[str] = []

    def translate_batch(
        self,
        *,
        texts: Sequence[str],
        source_language: str | None,
        target_language: str,
        format_hint: str,
        read_only_context: Sequence[str] = (),
        repair_candidates: Sequence[str] = (),
    ) -> TranslationBatchResult:
        """模拟正常与修复请求均失败。"""

        del texts, source_language, target_language, read_only_context, repair_candidates
        self.hints.append(format_hint)
        raise RuntimeError("synthetic provider failure")


class MissingCandidateTranslator:
    """返回空 batch, 用于模拟 provider 缺失全部 candidate。"""

    def __init__(self) -> None:
        """记录 pipeline 是否错误发起 repair 请求。"""

        self.hints: list[str] = []

    def translate_batch(
        self,
        *,
        texts: Sequence[str],
        source_language: str | None,
        target_language: str,
        format_hint: str,
        read_only_context: Sequence[str] = (),
        repair_candidates: Sequence[str] = (),
    ) -> TranslationBatchResult:
        """返回不含任何 index 的合法空 batch。"""

        del texts, source_language, target_language, read_only_context, repair_candidates
        self.hints.append(format_hint)
        return TranslationBatchResult(items=())


def test_pipeline_translates_body_tables_header_footer_without_overwriting_source(
    tmp_path: Path,
) -> None:
    """翻译正文、表格和页眉页脚, 同时保留格式与合并 cell。"""

    source = tmp_path / "source.docx"
    output = tmp_path / "translated.docx"
    document = Document()
    paragraph = document.add_paragraph()
    first = paragraph.add_run("Hello")
    first.bold = True
    second = paragraph.add_run(" world")
    second.italic = True
    table = document.add_table(rows=1, cols=3)
    merged = table.cell(0, 0).merge(table.cell(0, 1))
    merged.text = "Alpha"
    table.cell(0, 2).text = "123"
    document.sections[0].header.paragraphs[0].text = "Header text"
    document.sections[0].footer.paragraphs[0].text = "Footer text"
    document.save(source)
    original_bytes = source.read_bytes()

    replacements = {
        "Hello": "Bonjour",
        "world": "monde",
        "Alpha": "阿尔法",
        "Header text": "页眉",
        "Footer text": "页脚",
    }

    def transform(text: str) -> str:
        """只替换可见文本并保持 marker 与 tag 不变。"""

        for source_text, translated_text in replacements.items():
            text = text.replace(source_text, translated_text)
        return text

    translator = RecordingTranslator(transform)
    result = DocxPipeline(translator).translate_to(
        source_path=source,
        output_path=output,
        source_language="en",
        target_language="zh-CN",
    )

    translated = Document(output)
    assert source.read_bytes() == original_bytes
    assert translated.paragraphs[0].text == "Bonjour monde"
    assert translated.paragraphs[0].runs[0].bold is True
    assert translated.paragraphs[0].runs[1].italic is True
    assert translated.tables[0].cell(0, 0).text == "阿尔法"
    assert translated.tables[0].cell(0, 2).text == "123"
    assert translated.tables[0].cell(0, 0)._tc is translated.tables[0].cell(0, 1)._tc
    assert translated.sections[0].header.paragraphs[0].text == "页眉"
    assert translated.sections[0].footer.paragraphs[0].text == "页脚"
    chinese_runs = (
        translated.tables[0].cell(0, 0).paragraphs[0].runs[0],
        translated.sections[0].header.paragraphs[0].runs[0],
        translated.sections[0].footer.paragraphs[0].runs[0],
    )
    for run in chinese_runs:
        run_fonts = run._element.get_or_add_rPr().find(qn("w:rFonts"))
        assert run_fonts is not None
        assert run_fonts.get(qn("w:eastAsia")) == "Noto Sans SC"
        assert run_fonts.get(qn("w:ascii")) is None
        assert run_fonts.get(qn("w:hAnsi")) is None
        assert run_fonts.get(qn("w:cs")) is None
    assert result.translated_segments == 4
    assert result.fallback_segments == 0
    assert {hint for hint, _texts in translator.calls} == {"docx", "docx_table"}
    docx_batches = [texts for hint, texts in translator.calls if hint == "docx"]
    assert len(docx_batches) == 3
    assert docx_batches[0][0].startswith("[PARA_0]")
    assert docx_batches[1][0].startswith("[HEADER_0_PARA_0]")
    assert docx_batches[2][0].startswith("[FOOTER_0_PARA_0]")


def test_malformed_translation_gets_one_model_repair(tmp_path: Path) -> None:
    """首次结构损坏时只重试一次并记录模型修复 warning。"""

    source = _single_paragraph_document(tmp_path / "source.docx", "Original")
    output = tmp_path / "translated.docx"
    translator = RepairingTranslator()

    result = DocxPipeline(translator, translate_tables=False).translate_to(
        source_path=source,
        output_path=output,
        source_language="en",
        target_language="fr",
    )

    assert Document(output).paragraphs[0].text == "Translated"
    assert translator.hints == ["docx", "docx_repair"]
    assert translator.repair_candidates == [(), ("<span>broken",)]
    assert result.translated_segments == 1
    assert result.fallback_segments == 0
    assert result.warning_codes == ("docx_model_repair",)


def test_english_translation_adds_space_between_formatted_spans(tmp_path: Path) -> None:
    """英文译文跨格式 run 时应补词间空格且不改变 run 格式。"""

    source = tmp_path / "source.docx"
    output = tmp_path / "translated.docx"
    document = Document()
    paragraph = document.add_paragraph()
    first = paragraph.add_run("第一")
    first.bold = True
    second = paragraph.add_run("部分")
    second.italic = True
    document.save(source)

    DocxPipeline(
        RecordingTranslator(lambda text: text.replace("第一", "First").replace("部分", "part")),
        translate_tables=False,
    ).translate_to(
        source_path=source,
        output_path=output,
        source_language="zh-CN",
        target_language="en",
    )

    translated = Document(output).paragraphs[0]
    assert translated.text == "First part"
    assert translated.runs[0].text == "First "
    assert translated.runs[0].bold is True
    assert translated.runs[1].text == "part"
    assert translated.runs[1].italic is True


def test_english_translation_does_not_add_space_before_punctuation(tmp_path: Path) -> None:
    """English 标点位于后一 span 开头时不能被补出多余空格。"""

    source = tmp_path / "source.docx"
    output = tmp_path / "translated.docx"
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("你好")
    punctuated = paragraph.add_run("\uff0c世界")
    punctuated.italic = True
    document.save(source)

    DocxPipeline(
        RecordingTranslator(
            lambda text: text.replace("你好", "Hello").replace("\uff0c世界", ", world")
        ),
        translate_tables=False,
    ).translate_to(
        source_path=source,
        output_path=output,
        source_language="zh-CN",
        target_language="en-US",
    )

    assert Document(output).paragraphs[0].text == "Hello, world"


def test_literal_markup_is_escaped_then_unescaped_in_body_and_table(tmp_path: Path) -> None:
    """字面 XML-like 文本只作为正文翻译, 不能冒充 span 或 p 骨架。"""

    source = tmp_path / "source.docx"
    output = tmp_path / "translated.docx"
    body_text = "Body <span>tag</span> </span> <p> & <> Hello"
    table_text = "Cell <span>tag</span> </span> <p> & <> World"
    document = Document()
    document.add_paragraph(body_text)
    document.add_table(rows=1, cols=1).cell(0, 0).text = table_text
    document.save(source)
    translator = RecordingTranslator(
        lambda text: text.replace("Hello", "Bonjour").replace("World", "Monde")
    )

    result = DocxPipeline(translator).translate_to(
        source_path=source,
        output_path=output,
        source_language="en",
        target_language="fr",
    )

    translated = Document(output)
    assert translated.paragraphs[0].text == body_text.replace("Hello", "Bonjour")
    assert translated.tables[0].cell(0, 0).text == table_text.replace("World", "Monde")
    payloads = tuple(text for _hint, texts in translator.calls for text in texts)
    assert any(
        "Body &lt;span&gt;tag&lt;/span&gt; &lt;/span&gt; &lt;p&gt; &amp; &lt;&gt; Hello" in text
        for text in payloads
    )
    assert any(
        "Cell &lt;span&gt;tag&lt;/span&gt; &lt;/span&gt; &lt;p&gt; &amp; &lt;&gt; World" in text
        for text in payloads
    )
    assert result.translated_segments == 2
    assert result.fallback_segments == 0


def test_table_writeback_allows_newlines_between_cells(tmp_path: Path) -> None:
    """模型在相邻 cell tag 之间换行时仍应按原行骨架写回。"""

    source = tmp_path / "source.docx"
    output = tmp_path / "translated.docx"
    document = Document()
    table = document.add_table(rows=1, cols=2)
    first_cell = table.cell(0, 0)
    first_cell.text = ""
    first_cell.paragraphs[0].add_run("甲").bold = True
    first_cell.paragraphs[0].add_run("方").italic = True
    table.cell(0, 1).text = "乙"
    document.save(source)

    DocxPipeline(
        RecordingTranslator(
            lambda text: (
                text.replace("甲", "Alpha")
                .replace("方", "service")
                .replace("乙", "Beta")
                .replace("</span><span>", "</span>\n<span>")
                + "\n "
            )
        )
    ).translate_to(
        source_path=source,
        output_path=output,
        source_language="zh-CN",
        target_language="en",
    )

    translated = Document(output).tables[0]
    assert translated.cell(0, 0).text == "Alpha service"
    visible_runs = [run for run in translated.cell(0, 0).paragraphs[0].runs if run.text]
    assert visible_runs[0].bold is True
    assert visible_runs[1].italic is True
    assert translated.cell(0, 1).text == "Beta"


def test_table_writeback_allows_newlines_between_cell_paragraphs(tmp_path: Path) -> None:
    """多段 cell 的相邻 p tag 之间换行不能触发 markup mismatch。"""

    source = tmp_path / "source.docx"
    output = tmp_path / "translated.docx"
    document = Document()
    cell = document.add_table(rows=1, cols=1).cell(0, 0)
    cell.paragraphs[0].text = "第一段"
    cell.add_paragraph("第二段")
    document.save(source)

    DocxPipeline(
        RecordingTranslator(
            lambda text: (
                text.replace("第一段", "First paragraph")
                .replace("第二段", "Second paragraph")
                .replace("</p><p>", "</p>\n<p>")
            )
        )
    ).translate_to(
        source_path=source,
        output_path=output,
        source_language="zh-CN",
        target_language="en",
    )

    paragraphs = Document(output).tables[0].cell(0, 0).paragraphs
    assert [paragraph.text for paragraph in paragraphs] == [
        "First paragraph",
        "Second paragraph",
    ]


def test_provider_failure_returns_valid_source_text_fallback(tmp_path: Path) -> None:
    """整批调用失败时直接回退正文和表格, 且字面 tag 不会截断。"""

    source = tmp_path / "source.docx"
    output = tmp_path / "translated.docx"
    body_text = "Body </span> <p> & <> Original"
    table_text = "Table <span> </span> & <> Original"
    document = Document()
    document.add_paragraph(body_text)
    document.add_table(rows=1, cols=1).cell(0, 0).text = table_text
    document.save(source)
    translator = FailingTranslator()

    result = DocxPipeline(translator).translate_to(
        source_path=source,
        output_path=output,
        source_language=None,
        target_language="fr",
    )

    translated = Document(output)
    assert translated.paragraphs[0].text == body_text
    assert translated.tables[0].cell(0, 0).text == table_text
    assert translator.hints == ["docx", "docx_table"]
    assert result.translated_segments == 0
    assert result.fallback_segments == 2
    assert result.warning_codes == ("docx_segment_fallback", "docx_table_fallback")


def test_missing_candidate_falls_back_without_repair_call(tmp_path: Path) -> None:
    """provider 缺失 candidate 时直接 fallback, 不能制造逐 segment repair。"""

    source = _single_paragraph_document(tmp_path / "source.docx", "Original")
    output = tmp_path / "translated.docx"
    translator = MissingCandidateTranslator()

    result = DocxPipeline(translator, translate_tables=False).translate_to(
        source_path=source,
        output_path=output,
        source_language="en",
        target_language="fr",
    )

    assert Document(output).paragraphs[0].text == "Original"
    assert translator.hints == ["docx"]
    assert result.translated_segments == 0
    assert result.fallback_segments == 1
    assert result.warning_codes == ("docx_segment_fallback",)


def test_non_empty_spans_cannot_be_replaced_by_whitespace_around_tabs(tmp_path: Path) -> None:
    """结构合法的空白候选也不能清空由 tab 分隔的非空地址。"""

    source = tmp_path / "tabbed-address.docx"
    output = tmp_path / "translated.docx"
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("Huangxing Road Branch")
    for _index in range(3):
        paragraph.add_run().add_tab()
    paragraph.add_run("Gonghexin Road Branch")
    document.save(source)
    expected_text = Document(source).paragraphs[0].text

    def blank_visible_spans(text: str) -> str:
        """保留 marker 与 span 数量, 但故意把两个非空 slot 变为空格。"""

        marker = text[: text.index("]") + 1]
        return f"{marker}<span> </span><span> </span>"

    translator = RecordingTranslator(blank_visible_spans)
    result = DocxPipeline(translator, translate_tables=False).translate_to(
        source_path=source,
        output_path=output,
        source_language="en",
        target_language="zh-CN",
    )

    assert Document(output).paragraphs[0].text == expected_text
    assert [hint for hint, _texts in translator.calls] == ["docx", "docx_repair"]
    assert result.translated_segments == 0
    assert result.fallback_segments == 1
    assert result.warning_codes == ("docx_segment_fallback",)


def test_collapsed_multi_run_candidate_uses_joto_deterministic_repair(tmp_path: Path) -> None:
    """模型把多个 span 合成一个时保留完整译文并恢复源 run 数量。"""

    source = tmp_path / "formatted-address.docx"
    output = tmp_path / "translated.docx"
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("No. 1158, Jiuting Center Road, ").bold = True
    paragraph.add_run("Shuidian Road, Shanghai").italic = True
    document.save(source)

    def collapse_translation(text: str) -> str:
        """保留位置 marker, 但模拟 DeepSeek 把两个 span 合成一个。"""

        marker = text[: text.index("]") + 1]
        return f"{marker}<span>上海市水电路九亭中心路1158号</span>"

    translator = RecordingTranslator(collapse_translation)
    result = DocxPipeline(translator, translate_tables=False).translate_to(
        source_path=source,
        output_path=output,
        source_language="en",
        target_language="zh-CN",
    )

    translated = Document(output).paragraphs[0]
    assert translated.text == "上海市水电路九亭中心路1158号"
    assert len(translated.runs) == 2
    assert translated.runs[0].bold is True
    assert translated.runs[1].italic is True
    assert [hint for hint, _texts in translator.calls] == ["docx", "docx_repair"]
    assert result.translated_segments == 1
    assert result.fallback_segments == 0
    assert result.warning_codes == ("docx_deterministic_repair",)


def test_east_asian_compatible_source_font_is_preserved(tmp_path: Path) -> None:
    """w:hint=eastAsia 且字体落在 hAnsi 时仍识别为兼容中文。"""
    source = tmp_path / "east-asian-font.docx"
    output = tmp_path / "translated.docx"
    document = Document()
    run = document.add_paragraph().add_run("Original")
    run_properties = run._element.get_or_add_rPr()
    run_fonts = OxmlElement("w:rFonts")
    for attribute in ("ascii", "hAnsi", "cs"):
        run_fonts.set(qn(f"w:{attribute}"), "宋体")
    run_fonts.set(qn("w:hint"), "eastAsia")
    run_properties.insert(0, run_fonts)
    document.save(source)

    DocxPipeline(
        RecordingTranslator(lambda text: text.replace("Original", "中文")),
        translate_tables=False,
    ).translate_to(
        source_path=source,
        output_path=output,
        source_language="en",
        target_language="zh-CN",
    )

    translated_run = Document(output).paragraphs[0].runs[0]
    translated_fonts = translated_run._element.get_or_add_rPr().find(qn("w:rFonts"))
    assert translated_fonts.get(qn("w:eastAsia")) is None
    assert translated_fonts.get(qn("w:hAnsi")) == "宋体"
    assert translated_fonts.get(qn("w:hint")) == "eastAsia"


def test_latin_fallback_only_replaces_latin_font_slots(tmp_path: Path) -> None:
    """English fallback 同步 ascii/hAnsi, 但保留 East Asian 与 complex slot。"""

    source = tmp_path / "latin-font.docx"
    output = tmp_path / "translated.docx"
    document = Document()
    run = document.add_paragraph().add_run("原文")
    run_properties = run._element.get_or_add_rPr()
    run_fonts = OxmlElement("w:rFonts")
    for attribute, font_name in {
        "ascii": "Wingdings",
        "hAnsi": "Wingdings",
        "eastAsia": "SimSun",
        "cs": "Nirmala UI",
    }.items():
        run_fonts.set(qn(f"w:{attribute}"), font_name)
    run_properties.insert(0, run_fonts)
    document.save(source)

    DocxPipeline(
        RecordingTranslator(lambda text: text.replace("原文", "Contract terms")),
        translate_tables=False,
    ).translate_to(
        source_path=source,
        output_path=output,
        source_language="zh-CN",
        target_language="en",
    )

    translated_run = Document(output).paragraphs[0].runs[0]
    translated_fonts = translated_run._element.get_or_add_rPr().find(qn("w:rFonts"))
    assert translated_fonts.get(qn("w:ascii")) == "Calibri"
    assert translated_fonts.get(qn("w:hAnsi")) == "Calibri"
    assert translated_fonts.get(qn("w:eastAsia")) == "SimSun"
    assert translated_fonts.get(qn("w:cs")) == "Nirmala UI"


def test_structure_signature_still_detects_non_font_run_property_loss(tmp_path: Path) -> None:
    """忽略合法字体 fallback 时仍必须拦截 bold 与字号节点丢失。"""

    source = tmp_path / "formatted.docx"
    damaged = tmp_path / "damaged.docx"
    document = Document()
    run = document.add_paragraph().add_run("Original")
    run.bold = True
    run.font.size = Pt(22)
    document.save(source)

    damaged_document = Document(source)
    damaged_run = damaged_document.paragraphs[0].runs[0]
    damaged_run.bold = None
    damaged_run.font.size = None
    damaged_document.save(damaged)

    assert _package_signature(source) != _package_signature(damaged)


def test_validation_failure_keeps_existing_output_and_cleans_temp_file(
    tmp_path: Path,
) -> None:
    """临时 package 校验失败时不覆盖已有目标且清理临时文件。"""

    source = _single_paragraph_document(tmp_path / "source.docx", "Original")
    output = tmp_path / "translated.docx"
    output.write_bytes(b"existing-output")
    translator = RecordingTranslator(lambda text: text.replace("Original", "Translated"))

    class InvalidatingPipeline(DocxPipeline):
        """在原子替换前强制制造校验失败。"""

        def _validate_output(self, temporary_path, expected_signature) -> None:
            """模拟结构签名拒绝临时 package。"""

            del temporary_path, expected_signature
            raise ValueError("synthetic validation failure")

    with pytest.raises(ValueError, match="synthetic validation failure"):
        InvalidatingPipeline(translator, translate_tables=False).translate_to(
            source_path=source,
            output_path=output,
            source_language="en",
            target_language="fr",
        )

    assert output.read_bytes() == b"existing-output"
    assert not list(tmp_path.glob(".translated.docx.*.tmp.docx"))


def _single_paragraph_document(path: Path, text: str) -> Path:
    """创建只有一个正文段落的测试 DOCX。"""

    document = Document()
    document.add_paragraph(text)
    document.save(path)
    return path


def _request(source: Path, output_dir: Path) -> TranslationRequest:
    """构建共享 contract request 供输出命名测试复用。"""

    return TranslationRequest(
        source_path=source,
        output_dir=output_dir,
        source_language="en",
        target_language="zh-CN",
        provider_id="stub",
        model_id="stub",
    )


def test_shared_request_uses_safe_output_name(tmp_path: Path) -> None:
    """共享 request 路径使用目标语言后缀且绝不指向源文件。"""

    source = _single_paragraph_document(tmp_path / "source.docx", "Original")
    updates: list[TranslationProgress] = []
    result = DocxPipeline(
        RecordingTranslator(lambda text: text),
        translate_tables=False,
    ).translate(
        _request(source, tmp_path / "outputs"),
        report_progress=updates.append,
    )

    assert result.output_path.name == "source.zh-CN.docx"
    assert result.output_path != source
    assert updates[0] == TranslationProgress(stage="extracting")
    assert updates[1] == TranslationProgress(stage="translating", total_segments=1)
    assert updates[-1] == TranslationProgress(
        stage="formatting",
        processed_segments=1,
        total_segments=1,
    )


def test_bilingual_mode_derives_two_artifacts_without_translating_twice(tmp_path: Path) -> None:
    """双语版复用同一次翻译, 正文段内堆叠且表格只保留一份译文。"""

    source = tmp_path / "source.docx"
    document = Document()
    document.add_paragraph("第一条原文", style="List Number")
    document.add_paragraph("第二条原文", style="List Number")
    document.add_table(rows=1, cols=1).cell(0, 0).text = "表格原文"
    document.save(source)

    def transform(text: str) -> str:
        """保持 marker/span 骨架并替换测试正文。"""

        return (
            text.replace("第一条原文", "First item")
            .replace("第二条原文", "Second item")
            .replace("表格原文", "Table text")
        )

    translator = RecordingTranslator(transform)
    result = DocxPipeline(translator, bilingual=True).translate(
        _request(source, tmp_path / "outputs")
    )

    assert [artifact.kind for artifact in result.artifacts] == ["translated", "bilingual"]
    assert [hint for hint, _texts in translator.calls] == ["docx", "docx_table"]
    translated = Document(result.artifacts[0].path)
    bilingual = Document(result.artifacts[1].path)
    assert [paragraph.text for paragraph in translated.paragraphs] == [
        "First item",
        "Second item",
    ]
    assert [paragraph.text for paragraph in bilingual.paragraphs] == [
        "第一条原文\nFirst item",
        "第二条原文\nSecond item",
    ]
    assert len(bilingual.tables) == 1
    assert bilingual.tables[0].cell(0, 0).text == "Table text"
